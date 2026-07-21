"""Resume 文件解析 —— Sprint 5.8 (+ Sprint G 图片 OCR)。

定位:
- 把 PDF / docx / 图片文件二进制解析成纯文本, 交给候选人编辑后走旧 JSON 端点。
- PDF/docx 分支无副作用 (纯字节计算); 图片分支走 vision OCR (唯一的 LLM
  调用路径, Sprint G 起)。
- 任何拒绝路径 (大小 / mime / 解析后过短 / OCR 不可用) 一律抛 ResumeParseError,
  上游 api 路由映射成 422 + detail 给用户。

设计取舍:
- 库: pypdf (PDF) + python-docx (docx)。Resume 多单列纯文本, pdfplumber 的
  表格 / 多列优势用不上, 多依赖反而引 image 链路更重。
- 图片 (Sprint G): png/jpeg/webp 没文字层, 走 OpenAI vision (gpt-4o-mini)
  OCR。视觉模型能看懂排版, 比"扫描件 PDF 拍平文本"更准; 未配 key 时 vision
  返 stub, 本模块检测到就拒并提示贴文本。图片原图会发给 OpenAI (与文本简历
  出题同一数据流)。
- 校验: 5MB 上限远高于真实简历; mime + ext 双判防 mime 欺骗; 解析后文本
  < MIN_TEXT_CHARS 当作"扫描件图片 PDF 没文字层 / OCR 失败", 直接拒。
- 解析后做轻 normalize (合并空白行 / strip 行尾), 让前端 textarea 看着干净。
"""
from __future__ import annotations

import io
import re

from docx import Document  # python-docx
from pypdf import PdfReader

from src import llm

MAX_BYTES = 5 * 1024 * 1024              # 5 MB
MIN_TEXT_CHARS = 100                     # 解析后最少字数, 防扫描件图片 PDF / OCR 失败

_PDF_MIME = "application/pdf"
_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
# Sprint G: 图片 mime → 允许的扩展名 (gpt-4o vision 支持 png/jpeg/webp/非动 gif)
_IMAGE_MIMES: dict[str, tuple[str, ...]] = {
    "image/png": (".png",),
    "image/jpeg": (".jpg", ".jpeg"),
    "image/webp": (".webp",),
}
_PDF_EXTS = (".pdf",)
_DOCX_EXTS = (".docx",)

_OCR_TIMEOUT = 60.0        # 图片处理比纯文本慢, 给足; 前端有 loading 态
_OCR_SYSTEM = (
    "你是简历 OCR 工具。把图片里的简历内容完整、逐行转成纯文本, "
    "保留原始信息与顺序 (姓名、联系方式、教育经历、项目/实习/工作经历、"
    "技能、获奖等全部保留)。"
    "只输出简历文本本身: 不要任何解释, 不要 markdown 代码块包裹, "
    "绝不补充图片里没有的内容。"
)
_OCR_USER = "请把这张简历图片的内容转成纯文本。"


class ResumeParseError(ValueError):
    """解析任意一步失败 (大小 / 类型 / 内容不足 / 库报错) 统一抛这个。
    上游 api 路由把它映射成 HTTP 422 + str(e) 当 detail。"""


def parse_resume(*, filename: str, mime: str, blob: bytes) -> str:
    """主入口: 按 mime + extension dispatch, 返回 normalize 后的纯文本。

    抛 ResumeParseError 的场景:
    - blob 大小 > MAX_BYTES
    - mime / 扩展名不在白名单 OR 二者不匹配
    - 解析过程库报错 (PDF 损坏 / docx 不是合法 zip 等)
    - 解析出来文本 < MIN_TEXT_CHARS
    """
    if len(blob) > MAX_BYTES:
        raise ResumeParseError(
            f"文件超过 {MAX_BYTES // 1024 // 1024}MB 上限"
        )

    kind, mime_norm = _classify(filename, mime)
    try:
        if kind == "pdf":
            text = _parse_pdf(blob)
        elif kind == "docx":
            text = _parse_docx(blob)
        elif kind == "image":
            text = _ocr_image(blob, mime_norm)
        else:  # 防御性, 实际 _classify 不会返其他值
            raise ResumeParseError("不支持的文件类型")
    except ResumeParseError:
        raise
    except Exception as e:
        # pypdf.PdfReadError / docx.opc / openai 网络异常等统一吞成
        # ResumeParseError, 不向上抛底层栈, 候选人看到友好消息。
        raise ResumeParseError(f"解析失败: {type(e).__name__}") from e

    text = _normalize(text)
    if len(text) < MIN_TEXT_CHARS:
        raise ResumeParseError(
            f"解析后文本仅 {len(text)} 字符 (< {MIN_TEXT_CHARS}); "
            "可能是扫描件 PDF、图片过于模糊或排版图片化, 请直接粘贴文本"
        )
    return text


def _classify(filename: str, mime: str) -> tuple[str, str]:
    """按 mime 和 extension 双判, 二者必须指向同一类型 (防 mime 欺骗 +
    防扩展名欺骗)。返回 (kind, mime) —— kind ∈ 'pdf'|'docx'|'image'。"""
    name = filename.lower()

    if mime == _PDF_MIME and name.endswith(_PDF_EXTS):
        return "pdf", mime
    if mime == _DOCX_MIME and name.endswith(_DOCX_EXTS):
        return "docx", mime
    # Sprint G: 图片 mime + 对应扩展名匹配
    exts = _IMAGE_MIMES.get(mime)
    if exts and name.endswith(exts):
        return "image", mime

    # 任一不匹配, 拒绝。错误消息把支持的类型列出来, 用户能立刻看出问题。
    raise ResumeParseError(
        f"文件类型不被支持: filename={filename!r} mime={mime!r}; "
        "仅接受 PDF (.pdf)、docx (.docx) 和图片 (.png/.jpg/.webp)"
    )


def _ocr_image(blob: bytes, mime: str) -> str:
    """图片简历 → vision OCR → 纯文本 (Sprint G)。
    未配 vision (返 stub) / 空输出 → 抛 ResumeParseError, 让用户改贴文本。
    vision 网络异常由上层 parse_resume 的 except 兜成友好错误。"""
    text = llm.complete_vision(
        _OCR_SYSTEM, _OCR_USER, [(mime, blob)], timeout=_OCR_TIMEOUT,
    )
    if not text or llm.is_stub(text):
        raise ResumeParseError(
            "图片识别当前不可用 (未配置视觉模型); "
            "请上传 PDF / docx, 或直接粘贴简历文本"
        )
    return text


def _parse_pdf(blob: bytes) -> str:
    """逐页提取文本。pypdf 6.x 的 page.extract_text() 偶尔返 None, 兜底空串。"""
    reader = PdfReader(io.BytesIO(blob))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _parse_docx(blob: bytes) -> str:
    """docx 主体: paragraphs + tables. Resume 常用表格塞经历, 不能漏。"""
    doc = Document(io.BytesIO(blob))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


_BLANK_LINES_RE = re.compile(r"\n{3,}")
_TRAILING_WS_RE = re.compile(r"[ \t]+\n")


def _normalize(text: str) -> str:
    """轻 normalize: strip 行尾空白; 把 >=3 连续换行压成 2 (保留段落分隔但
    不留大片空白)。前端 textarea 显示干净, HR / Planner 看也舒服。"""
    text = _TRAILING_WS_RE.sub("\n", text)
    text = _BLANK_LINES_RE.sub("\n\n", text)
    return text.strip()
